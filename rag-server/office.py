from dotenv import load_dotenv
import os  # 파일 경로 구성, 디렉토리 탐색 등 운영체제 관련 기능을 사용할 수 있는 표준 모듈
load_dotenv()
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")   # OpenAI GPT 모델을 사용하기 위한 API 키를 환경 변수로 등록 (주의: 실제 키로 대체 필요)

import logging  # 로그 출력을 위한 표준 파이썬 모듈
from mcp.server.fastmcp import FastMCP  # MCP 서버를 간편하게 생성하는 클래스 (도구 등록 및 실행 가능)

# LangChain 관련 모듈 (문서 처리 및 검색 기반 QA 기능 구현)
from langchain.text_splitter import RecursiveCharacterTextSplitter  # 문서를 일정 단위로 나누는 유틸리티
from langchain_community.vectorstores import Chroma  # 문서를 벡터로 저장하고 검색 가능한 DB (Chroma)
from langchain.schema import Document  # LangChain에서 사용하는 문서 객체 타입
from langchain.chains import RetrievalQA  # 검색 기반 질문 응답(RAG) 체인 생성 도구
from langchain_openai import ChatOpenAI, OpenAIEmbeddings  # OpenAI의 GPT 모델 및 임베딩 모델을 사용하는 LangChain 래퍼

# Word 및 Excel 문서 처리용 모듈
from docx import Document as WordDocument  # Word 문서를 불러오기 위한 클래스 (이름 충돌 피하기 위해 별칭 사용)
import pandas as pd  # Excel 파일을 불러오기 위한 데이터 처리 라이브러리

# 로그 레벨을 INFO 이상으로 설정 (로그를 콘솔에 출력)
logging.basicConfig(level=logging.INFO)

# MCP 서버 인스턴스를 생성 ("Office-RAG"이라는 이름으로 구분)
mcp = FastMCP("Office-RAG")

# 문서들이 저장된 폴더 경로를 지정 (여기서는 로컬 PC 경로)
OFFICE_DIR = "C:/Users/admin/test-server/rag-server/office"

# Word(.docx) 및 Excel(.xlsx) 파일을 읽어 LangChain 문서 리스트로 변환하는 함수
def load_office_documents(folder_path: str) -> list[Document]:
    docs = []  # 최종 결과로 반환할 문서 리스트

    # 지정된 폴더 내 모든 파일에 대해 반복
    for filename in os.listdir(folder_path):
        path = os.path.join(folder_path, filename)

        # Word 문서인 경우
        if filename.endswith(".docx"):
            word = WordDocument(path)  # Word 문서 로드
            # 각 문단의 텍스트를 모아 하나의 문자열로 결합 (빈 줄 제외)
            full_text = "\n".join([p.text for p in word.paragraphs if p.text.strip()])
            # LangChain 문서 객체로 추가 (source에 파일 이름 저장)
            docs.append(Document(page_content=full_text, metadata={"source": filename}))

        # Excel 문서인 경우
        elif filename.endswith(".xlsx"):
            try:
                excel = pd.read_excel(path, sheet_name=None)  # 모든 시트 로드 (dict로 반환)
                for sheet_name, df in excel.items():
                    # 시트 데이터를 문자열로 변환하여 텍스트화
                    text = df.astype(str).to_string(index=False)
                    # 각 시트를 하나의 문서로 저장 (파일명 + 시트명 기록)
                    docs.append(Document(page_content=text, metadata={"source": f"{filename} - {sheet_name}"}))
            except Exception as e:
                # Excel 파일 로딩 중 오류 발생 시 로그에 기록
                logging.error(f"엑셀 파일 처리 오류: {filename} - {e}")

    return docs  # Word + Excel 문서를 모두 포함한 리스트 반환

# 문서 로딩 및 분할
raw_docs = load_office_documents(OFFICE_DIR)  # 폴더 내 Word/Excel 문서 불러오기
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)  # 한 문서를 500자 단위로 나누고 50자 중첩
docs = splitter.split_documents(raw_docs)  # 분할된 문서 리스트 생성

# 임베딩 모델 및 LLM(GPT) 모델 초기화
embeddings = OpenAIEmbeddings()  # 텍스트를 벡터로 변환하는 OpenAI 임베딩
llm = ChatOpenAI(model="gpt-5")  # GPT-5 모델 사용 (대화형 응답 생성)

# 문서를 벡터 DB에 저장 
vectorstore = Chroma.from_documents(
    documents=docs,  # 분할된 문서 리스트
    embedding=embeddings  # 텍스트를 벡터로 변환할 임베딩 모델
)

# 검색 기반 질문 응답 체인(RAG) 구성
qa_chain = RetrievalQA.from_chain_type(
    llm=llm,  # 사용할 GPT 모델
    retriever=vectorstore.as_retriever()  # 벡터 DB에서 문서를 검색할 리트리버
)

# MCP 도구로 등록 (LLM이 자동으로 호출할 수 있게 함)
@mcp.tool()
def ask_office(query: str) -> str:
    """폴더 내 Word/Excel 문서를 기반으로 질문에 답변합니다."""
    logging.info(f"Query: {query}")
    try:
        return qa_chain.run(query)
    except Exception as e:
        logging.error(f"ask_office 실행 중 오류: {e}")
        return (
            "죄송합니다. ask_office 도구 실행 중 오류가 발생했습니다.\n"
            f"[에러 내용] {e}"
        )

# MCP 서버 실행 (표준 입출력 기반 - Cursor, Claude Desktop 등과 연동 가능)
if __name__ == "__main__":
    mcp.run(transport="stdio")