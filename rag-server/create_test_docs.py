#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
테스트용 Word 및 Excel 문서 생성 스크립트
"""

from docx import Document
import pandas as pd
import os

def create_test_word_document():
    """테스트용 Word 문서 생성"""
    doc = Document()
    
    # 제목 추가
    doc.add_heading('MCP 서버와 RAG 시스템 테스트 문서', 0)
    
    # 소개 문단
    doc.add_paragraph('이 문서는 MCP 서버와 RAG 시스템의 테스트를 위해 생성되었습니다.')
    
    # 주요 내용
    doc.add_heading('주요 기능', level=1)
    doc.add_paragraph('• MCP 서버를 통한 도구 등록 및 관리')
    doc.add_paragraph('• LangChain을 활용한 RAG 시스템 구축')
    doc.add_paragraph('• PDF 문서 기반 질문응답 시스템 구현')
    doc.add_paragraph('• AI 모델과 외부 도구의 연동')
    
    # 기술 스택
    doc.add_heading('사용 기술', level=1)
    doc.add_paragraph('• Python 3.13')
    doc.add_paragraph('• LangChain 프레임워크')
    doc.add_paragraph('• OpenAI API')
    doc.add_paragraph('• Chroma 벡터 데이터베이스')
    doc.add_paragraph('• MCP 프로토콜')
    
    # 결론
    doc.add_heading('결론', level=1)
    doc.add_paragraph('이번 실습을 통해 MCP 서버의 기본 구조와 RAG 시스템의 구현 방법을 학습했습니다.')
    
    # 파일 저장
    filename = '테스트_문서.docx'
    doc.save(filename)
    print(f"✅ Word 문서 생성 완료: {filename}")
    return filename

def create_test_excel_document():
    """테스트용 Excel 문서 생성"""
    # 샘플 데이터
    data = {
        '기능명': ['MCP 서버', 'RAG 시스템', 'LangChain', 'OpenAI API', 'Chroma DB'],
        '설명': [
            'AI와 도구를 연결하는 프로토콜',
            '검색과 생성을 결합한 시스템',
            'LLM 애플리케이션 개발 프레임워크',
            'OpenAI의 언어 모델 API',
            '벡터 데이터베이스'
        ],
        '상태': ['완료', '완료', '완료', '완료', '완료'],
        '우선순위': [1, 1, 2, 2, 3]
    }
    
    # DataFrame 생성
    df = pd.DataFrame(data)
    
    # Excel 파일로 저장
    filename = '테스트_데이터.xlsx'
    with pd.ExcelWriter(filename, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='기능 목록', index=False)
        
        # 추가 시트 생성
        summary_data = {
            '항목': ['총 기능 수', '완료된 기능', '진행 중인 기능'],
            '수치': [len(data['기능명']), len(data['기능명']), 0]
        }
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, sheet_name='요약', index=False)
    
    print(f"✅ Excel 문서 생성 완료: {filename}")
    return filename

def main():
    """메인 함수"""
    print("🚀 테스트용 문서 생성 시작...")
    
    try:
        # Word 문서 생성
        word_file = create_test_word_document()
        
        # Excel 문서 생성
        excel_file = create_test_excel_document()
        
        print("\n🎉 모든 테스트 문서 생성 완료!")
        print(f"📄 Word 문서: {word_file}")
        print(f"📊 Excel 문서: {excel_file}")
        print("\n💡 이제 office.py를 실행하여 테스트할 수 있습니다.")
        
    except Exception as e:
        print(f"❌ 오류 발생: {e}")
        print("💡 필요한 패키지가 설치되어 있는지 확인하세요:")
        print("   pip install python-docx pandas openpyxl")

if __name__ == "__main__":
    main()
